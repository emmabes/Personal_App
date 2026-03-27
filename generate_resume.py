"""
generate_resume.py
------------------
Generates Erik_Mabes_Resume.docx (and optionally Erik_Mabes_Resume.pdf)
in the project root.

Design language:
  - Light background (white) — readable everywhere
  - Dark teal   #004F5E  → name, section headers
  - Amber-orange #B84700  → dates, role labels (readable on white)
  - Near-black   #1A1A1A  → body bullets
  - Thin cyan-tinted rule under each section header (#00B4C8)
  - Calibri throughout — safe for all major ATS systems

ATS rules enforced:
  - Single column, no tables, no text boxes
  - Contact block in document body (not Word header/footer)
  - Standard bullet character (•)
  - Dates on same line as company / title
  - .docx saved cleanly

Run:
    python generate_resume.py          # docx only
    python generate_resume.py --pdf    # docx + PDF (requires MS Word on Windows)
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette (all readable on white) ──────────────────────────────────────────
TEAL      = RGBColor(0x00, 0x4F, 0x5E)   # name, section headers
ORANGE    = RGBColor(0xB8, 0x47, 0x00)   # dates, role labels, company
BODY      = RGBColor(0x1A, 0x1A, 0x1A)   # bullet text
RULE_RGB  = RGBColor(0x00, 0xB4, 0xC8)   # thin rule under section headers
MUTED     = RGBColor(0x55, 0x55, 0x55)   # contact line, stack line

FONT_NAME = 'Calibri'

# ── Data ─────────────────────────────────────────────────────────────────────
contact = {
    'name':     'Erik Mabes',
    'location': 'Phoenix, AZ 85048',
    'phone':    '928-607-9156',
    'email':    'erik.a.mabes@outlook.com',
    'links': [
        'linkedin.com/in/erik-mabes-4b6133184',
        'github.com/emmabes',
        'erikmabes.com',
    ],
}

summary = (
    'Java and AWS Software Engineer with production experience across the full SDLC at '
    'Amazon-scale. Spearheaded an image inspection feature that reduced upload errors by 9% '
    'and saved ~$50M annually; collaborated on a listings error initiative saving ~$250M '
    'annually. Coded against 16+ AWS services via IaC in a high-availability, EC2-based '
    'microservices environment. Cross-functional background spanning big-tech engineering, '
    'Lean Six Sigma project management, and operational leadership — enabling rapid '
    'identification of high-value technology opportunities across organizations of any size.'
)

skills = [
    ('Languages',          'Java, Python, TypeScript, JavaScript, SQL, VBA'),
    ('Frontend',           'React, TypeScript, HTML/CSS'),
    ('Backend',            'Java, REST APIs, Microservices, Distributed Systems, Event-Driven Architecture, SQS'),
    ('Cloud / AWS',        'API Gateway, Lambda, EC2, DynamoDB, S3, IAM, Secrets Manager, CloudWatch, '
                           'CodePipeline, CodeBuild, CodeCommit, KMS, Amplify, CloudFront, Glue, SQS, Polly'),
    ('DevOps / CI/CD',     'Git, Docker, CI/CD Pipelines, Infrastructure as Code (IaC), On-Call Operations'),
    ('Testing',            'JUnit, Postman, Test-Driven Development, Automated Testing, Performance Testing'),
    ('Data / Analytics',   'Python Scripting, Data Pipeline Design, Power Query, Excel, VBA Macros'),
    ('Project Management', 'Agile, Scrum, Kanban, Lean Six Sigma'),
]

jobs = [
    {
        'company': 'TKM Restaurants, Inc.',
        'role':    'Director of Operations',
        'dates':   'Feb 2025 – Present',
        'bullets': [
            'Built Python-based report consolidators, generators, payroll automators, and data mapping tools actively used in daily business operations',
            'Integrated with Silverware POS APIs, vendor APIs, and automated reporting systems to improve managerial and administrative workflows',
            'Developed a central KPI dashboard using Python data collection/processing, Excel, and Power Query',
            'Lead 63 employees including 4 managers across a high-volume service operation; oversee hiring, training, scheduling, and operational strategy',
        ],
    },
    {
        'company': 'TKM Restaurants, Inc.',
        'role':    'Project Manager',
        'dates':   'Feb 2024 – Jan 2025',
        'bullets': [
            'Revived a remodel stalled for 7+ months; compressed an 8-week projected shutdown to 3.5 weeks using Lean Six Sigma strategies',
            'Recouped approximately $360,000 in sales from the shortened closure; managed timelines, budgets, contractors, and compliance',
            'Personally executed plumbing, electrical, and IT networking tasks to eliminate blockers in a resource-constrained environment',
        ],
    },
    {
        'company': 'Amazon.com',
        'role':    'Software Development Engineer I',
        'dates':   'Oct 2022 – Dec 2023',
        'bullets': [
            'Spearheaded an image URL inspection feature for 3rd-party seller bulk upload experience, reducing upload errors by 9% and saving ~$50M in annual operational costs',
            'Collaborated on a team initiative that reduced listing error rates from 60%+ to 14%, contributing to ~$250M in annual cost savings',
            'Contributed to a 61% reduction in seller report generation times through backend microservice improvements',
            'Owned full SDLC responsibilities: CI/CD pipeline management, on-call production incident response, code review, automated and performance testing',
            'Designed and implemented asynchronous system calls and revised data processing hierarchies across team-owned microservices',
            'Unified 3rd-party seller and 1st-party vendor bulk upload pipelines into a shared Java backend system across team boundaries',
            'AWS services via IaC: API Gateway, Lambda, EC2, DynamoDB, S3, IAM, Secrets Manager, CloudWatch, CodePipeline, CodeBuild, CodeCommit, Polly, KMS, Amplify, CloudFront, Glue, SQS',
        ],
    },
    {
        'company': 'Amazon Technical Academy',
        'role':    'SDE Training Program',
        'dates':   'Feb 2022 – Oct 2022',
        'bullets': [
            'Completed a 9-month, paid-leave internal SDE program: 12% acceptance rate, 80% graduation rate, fewer than 4% of graduates placed directly to SDE I',
            'Curriculum developed by Amazon Software Directors, Managers, and Principal Engineers; covered Java, AWS, microservices, CI/CD, Docker, Agile/Scrum, and TDD',
            'Achieved 99% overall program grade; led team to Top 3 of 44 teams in capstone web application project',
            'Directly promoted to SDE I, bypassing the standard Junior Developer / Apprenticeship track',
        ],
    },
    {
        'company': 'Amazon.com',
        'role':    'Project Coordinator',
        'dates':   'Feb 2021 – Feb 2022',
        'bullets': [
            'Primary owner of the WFS Launch Tracker — a C-Suite data source synthesizing 32,000+ data points across 523 active launch sites in 2021',
            'Reduced Launch Tracker time demands by 42% through Python and VBA automation',
            'Reduced overall manual workflows by 62% by converting email-based reporting into an automated web dashboard',
        ],
    },
    {
        'company': 'Amazon.com',
        'role':    'Staffing Ambassador',
        'dates':   'Mar 2020 – Feb 2021',
        'bullets': [
            'Self-assigned to an internal Optimization Team; reduced key report generation times from 1.5+ hours to under 10 minutes each',
            'Reduced team weekly hours from ~120 to under 60; automation saved $30,000 annually per field team, $210,000 annually after rollout to additional teams',
            'Supported record launch of 13 sites in 4 months through process optimization and auditing',
        ],
    },
]

projects = [
    {
        'name':   'erikmabes.com — Personal Website',
        'status': 'Live',
        'stack':  'AWS (CDK, Lambda, CloudFront, Cognito, API Gateway, S3), Python, React, TypeScript CDK',
        'bullets': [
            'Rate-limited resume download backend (Python Lambda); browser-based games; interactive animated resume visualizer',
            'User authentication with RBAC via Cognito; fully deployed on AWS with active CI/CD pipelines',
        ],
    },
    {
        'name':   'Mobile Application — Meal Planning & Household Management',
        'status': 'In Development',
        'stack':  'React Native, Node.js',
        'bullets': [
            'Meal selection and planning, automated shopping list generation, inventory tracking, recipe database, cross-service price comparisons',
        ],
    },
    {
        'name':   'TKM Operational Tooling',
        'status': 'Production',
        'stack':  'Python, Silverware POS API, Excel, VBA, Power Query',
        'bullets': [
            'Production tools for report generation, payroll automation, and KPI tracking actively used by the business operations team',
        ],
    },
]

education = [
    {
        'school': 'University of Arizona',
        'degree': 'B.S. Neuroscience & Cognitive Sciences',
        'note':   'Neurobiology emphasis · Psychology minor',
        'date':   'Dec 2017',
    },
    {
        'school': 'Amazon Technical Academy',
        'degree': 'Internal SDE Program Certificate',
        'note':   '9-month program · 12% acceptance rate · 99% overall grade · Direct-to-SDE I placement',
        'date':   'Oct 2022',
    },
    {
        'school': 'Coconino Community College',
        'degree': 'A.A.S. Applied Sciences of Pre-Health',
        'note':   'Cum Laude',
        'date':   'May 2014',
    },
]


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_paragraph(doc, text='', style='Normal', space_before=0, space_after=4,
                  left_indent=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    if text:
        p.add_run(text)
    return p


def add_bottom_border(paragraph, color_hex='00B4C8', thickness=6):
    """Adds a thin coloured bottom border to a paragraph (ATS-safe: it's a paragraph border)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    str(thickness))   # 1/8 pt units; 6 = 0.75pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_header(doc, title):
    """ALL-CAPS teal section header with a thin cyan rule underneath."""
    p = add_paragraph(doc, space_before=10, space_after=2)
    run = p.add_run(title.upper())
    set_font(run, 10, bold=True, color=TEAL)
    add_bottom_border(p)
    return p


def bullet(doc, text, indent=0.2):
    p = add_paragraph(doc, space_before=0, space_after=2, left_indent=indent)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    run = p.add_run(f'\u2022  {text}')
    set_font(run, 10, color=BODY)
    return p


def job_header(doc, company, role, dates):
    """Company (bold teal) | Role (italic, muted) right-aligned dates (orange)."""
    p = add_paragraph(doc, space_before=6, space_after=1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

    r1 = p.add_run(company)
    set_font(r1, 10.5, bold=True, color=TEAL)

    r2 = p.add_run('  ·  ')
    set_font(r2, 10, color=MUTED)

    r3 = p.add_run(role)
    set_font(r3, 10, italic=True, color=MUTED)

    r4 = p.add_run(f'\t{dates}')
    set_font(r4, 9.5, bold=True, color=ORANGE)

    return p


def edu_entry(doc, school, degree, note, date):
    p = add_paragraph(doc, space_before=5, space_after=1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)

    r1 = p.add_run(school)
    set_font(r1, 10.5, bold=True, color=TEAL)

    r4 = p.add_run(f'\t{date}')
    set_font(r4, 9.5, bold=True, color=ORANGE)

    p2 = add_paragraph(doc, space_before=0, space_after=0)
    r5 = p2.add_run(degree)
    set_font(r5, 10, color=BODY)

    p3 = add_paragraph(doc, space_before=0, space_after=2)
    r6 = p3.add_run(note)
    set_font(r6, 9.5, italic=True, color=MUTED)


def project_header(doc, name, status, stack):
    p = add_paragraph(doc, space_before=6, space_after=1)
    r1 = p.add_run(name)
    set_font(r1, 10.5, bold=True, color=TEAL)
    r2 = p.add_run(f'  ({status})')
    set_font(r2, 9.5, color=ORANGE)

    p2 = add_paragraph(doc, space_before=0, space_after=2)
    r3 = p2.add_run(f'Stack: {stack}')
    set_font(r3, 9.5, italic=True, color=MUTED)


# ── Build document ────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins: 0.75" all around — gives ~7" text width
    for section in doc.sections:
        section.top_margin    = Cm(1.9)
        section.bottom_margin = Cm(1.9)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # ── NAME ──────────────────────────────────────────────────────────────────
    p_name = add_paragraph(doc, space_before=0, space_after=2)
    p_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p_name.add_run(contact['name'].upper())
    r.font.name   = FONT_NAME
    r.font.size   = Pt(26)
    r.font.bold   = True
    r.font.color.rgb = TEAL

    # ── CONTACT LINE ─────────────────────────────────────────────────────────
    contact_parts = [contact['location'], contact['phone'], contact['email']] + contact['links']
    p_contact = add_paragraph(doc, space_before=0, space_after=8)
    add_bottom_border(p_contact, color_hex='00B4C8', thickness=4)
    r_c = p_contact.add_run('  ·  '.join(contact_parts))
    set_font(r_c, 9.5, color=MUTED)

    # ── SUMMARY ───────────────────────────────────────────────────────────────
    section_header(doc, 'Professional Summary')
    p_sum = add_paragraph(doc, space_before=3, space_after=0)
    r_sum = p_sum.add_run(summary)
    set_font(r_sum, 10, color=BODY)
    p_sum.paragraph_format.line_spacing = Pt(14)

    # ── SKILLS ────────────────────────────────────────────────────────────────
    section_header(doc, 'Technical Skills')
    for label, value in skills:
        p_sk = add_paragraph(doc, space_before=1, space_after=1)
        r_lbl = p_sk.add_run(f'{label}:  ')
        set_font(r_lbl, 10, bold=True, color=ORANGE)
        r_val = p_sk.add_run(value)
        set_font(r_val, 10, color=BODY)

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    section_header(doc, 'Work Experience')
    for job in jobs:
        job_header(doc, job['company'], job['role'], job['dates'])
        for b in job['bullets']:
            bullet(doc, b)

    # ── PROJECTS ─────────────────────────────────────────────────────────────
    section_header(doc, 'Projects')
    for proj in projects:
        project_header(doc, proj['name'], proj['status'], proj['stack'])
        for b in proj['bullets']:
            bullet(doc, b)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    section_header(doc, 'Education')
    for e in education:
        edu_entry(doc, e['school'], e['degree'], e['note'], e['date'])

    doc.save('Erik_Mabes_Resume.docx')
    print('OK  Erik_Mabes_Resume.docx written.')
    return doc


# ── PDF conversion ────────────────────────────────────────────────────────────

def to_pdf():
    """
    Converts via Microsoft Word COM automation (Windows + Word required).
    Falls back to a helpful error message if Word is not available.
    """
    import subprocess, os
    docx_path = os.path.abspath('Erik_Mabes_Resume.docx')
    pdf_path  = os.path.abspath('Erik_Mabes_Resume.pdf')
    ps_script = (
        f"$word = New-Object -ComObject Word.Application; "
        f"$word.Visible = $false; "
        f"$doc = $word.Documents.Open('{docx_path}'); "
        f"$doc.SaveAs('{pdf_path}', 17); "
        f"$doc.Close(); "
        f"$word.Quit()"
    )
    result = subprocess.run(
        ['powershell', '-Command', ps_script],
        capture_output=True, text=True
    )
    if result.returncode == 0 and os.path.exists(pdf_path):
        print('OK  Erik_Mabes_Resume.pdf written.')
    else:
        print('FAIL  PDF conversion failed.')
        print('   Ensure Microsoft Word is installed.')
        if result.stderr:
            print(result.stderr[:400])


if __name__ == '__main__':
    build()
    if '--pdf' in sys.argv:
        to_pdf()
